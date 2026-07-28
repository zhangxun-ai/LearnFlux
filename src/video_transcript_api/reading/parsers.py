"""Small, synchronous parsers used by the first reading MVP."""

from __future__ import annotations

import html
import base64
import binascii
import importlib.util
import json
import mimetypes
import posixpath
import re
import subprocess
import sys
import tempfile
import zipfile
from concurrent.futures import Future, ThreadPoolExecutor
from dataclasses import dataclass, field, replace
from pathlib import Path
from xml.etree import ElementTree

import markdown
import nh3
from bs4 import BeautifulSoup

from .assets import ParsedAsset, ReadingAssetError


class LocalOcrUnavailable(RuntimeError):
    """The optional local OCR runtime is not installed."""


_LOCAL_OCR_MAX_RENDER_SIDE = 3200
_LOCAL_OCR_PREFERRED_RENDER_SCALE = 2.0
_LOCAL_OCR_RUNNER_TIMEOUT_SECONDS = 180
_LOCAL_OCR_MAX_WORKERS = 2


@dataclass(frozen=True)
class ParsedChapter:
    title: str
    plain_text: str
    sanitized_html: str
    source_page: int | None = None
    chapter_key: str = ""
    level: int = 1
    parent_key: str | None = None
    epub_cfi: str = ""


@dataclass(frozen=True)
class ParsedOutlineItem:
    chapter_key: str
    title: str
    level: int
    parent_key: str | None = None
    source_page: int | None = None
    epub_cfi: str = ""


@dataclass
class ParsedDocument:
    chapters: list[ParsedChapter]
    outline: list[ParsedOutlineItem]
    assets: list[ParsedAsset]
    pages: list[StructuredPage] = field(default_factory=list)

    def __bool__(self) -> bool:
        return bool(self.chapters)


@dataclass(frozen=True)
class OcrBlock:
    """One recognized text block with optional page-relative geometry."""

    text: str
    bbox: tuple[float, float, float, float] | None
    confidence: float | None
    kind: str = "unknown"


@dataclass(frozen=True)
class PageQuality:
    """Local confidence assessment for one structured PDF page."""

    status: str
    score: float
    reasons: tuple[str, ...] = ()


@dataclass(frozen=True)
class StructuredPage:
    """Page-level OCR/native text result before document assembly."""

    source_page: int
    width: float
    height: float
    extraction_mode: str
    blocks: tuple[OcrBlock, ...] | list[OcrBlock]
    body_blocks: tuple[OcrBlock, ...] = ()
    footnote_blocks: tuple[OcrBlock, ...] = ()


def local_ocr_render_scale(page_width: float, page_height: float) -> float:
    """Fit PDF rasterization to Paddle's detector limit before OCR starts."""
    largest_side = max(float(page_width), float(page_height), 1.0)
    return min(_LOCAL_OCR_PREFERRED_RENDER_SCALE, _LOCAL_OCR_MAX_RENDER_SIDE / largest_side)


def local_ocr_worker_count(page_count: int) -> int:
    """Keep local OCR parallelism within the verified native-runtime limit."""
    return max(1, min(_LOCAL_OCR_MAX_WORKERS, int(page_count)))


def _block_top(block: OcrBlock) -> float:
    return block.bbox[1] if block.bbox is not None else float("inf")


def _block_center_x(block: OcrBlock) -> float:
    if block.bbox is None:
        return float("inf")
    return (block.bbox[0] + block.bbox[2]) / 2


def _is_bottom_footnote(page: StructuredPage, block: OcrBlock) -> bool:
    if block.bbox is None or page.height <= 0:
        return False
    top = block.bbox[1]
    height = max(0.0, block.bbox[3] - block.bbox[1])
    return top >= page.height * 0.9 and height <= page.height * 0.04


def _sort_body_blocks(page: StructuredPage, blocks: list[OcrBlock]) -> list[OcrBlock]:
    if len(blocks) < 2 or page.width <= 0:
        return sorted(blocks, key=lambda item: (_block_top(item), _block_center_x(item)))
    centers = [_block_center_x(block) for block in blocks]
    if max(centers) - min(centers) < page.width * 0.3:
        return sorted(blocks, key=lambda item: (_block_top(item), _block_center_x(item)))
    split = (max(centers) + min(centers)) / 2
    left = [block for block in blocks if _block_center_x(block) <= split]
    right = [block for block in blocks if _block_center_x(block) > split]
    return sorted(left, key=_block_top) + sorted(right, key=_block_top)


def classify_page_blocks(page: StructuredPage) -> StructuredPage:
    """Separate footer notes and establish a deterministic local reading order."""
    body = [
        block
        for block in page.blocks
        if block.text.strip() and not _is_bottom_footnote(page, block)
    ]
    footnotes = [
        block
        for block in page.blocks
        if block.text.strip() and _is_bottom_footnote(page, block)
    ]
    return replace(
        page,
        body_blocks=tuple(_sort_body_blocks(page, body)),
        footnote_blocks=tuple(sorted(footnotes, key=_block_top)),
    )


def classify_document_pages(pages: list[StructuredPage]) -> list[StructuredPage]:
    """Remove text repeated in page margins after individual page ordering."""
    classified = [classify_page_blocks(page) for page in pages]
    candidates: dict[str, set[int]] = {}
    for page in classified:
        for block in page.body_blocks:
            if block.bbox is None or page.height <= 0:
                continue
            top, bottom = block.bbox[1], block.bbox[3]
            if top > page.height * 0.1 and bottom < page.height * 0.95:
                continue
            signature = re.sub(r"\s+", "", block.text).strip()
            if signature:
                candidates.setdefault(signature, set()).add(page.source_page)
    repeated = {signature for signature, source_pages in candidates.items() if len(source_pages) >= 2}
    if not repeated:
        return classified
    return [
        replace(
            page,
            body_blocks=tuple(
                block
                for block in page.body_blocks
                if re.sub(r"\s+", "", block.text).strip() not in repeated
            ),
        )
        for page in classified
    ]


def score_page_quality(page: StructuredPage) -> PageQuality:
    """Classify local page output without claiming unsupported OCR quality."""
    blocks = [block for block in page.blocks if block.text.strip()]
    if not blocks:
        return PageQuality("attention", 0.0, ("no_text_blocks",))
    if page.extraction_mode == "local_ocr" and any(
        block.bbox is None for block in blocks
    ):
        return PageQuality("attention", 0.0, ("missing_coordinates",))
    if page.extraction_mode == "native_text":
        if any(block.bbox is None for block in blocks):
            return PageQuality("attention", 0.0, ("missing_coordinates",))
        return PageQuality("good", 1.0)
    confidences = [block.confidence for block in blocks if block.confidence is not None]
    if not confidences:
        return PageQuality("warning", 0.5, ("missing_confidence",))
    score = sum(confidences) / len(confidences)
    if score < 0.7:
        return PageQuality("attention", score, ("low_confidence",))
    if score < 0.9:
        return PageQuality("warning", score, ("reduced_confidence",))
    return PageQuality("good", score)


def _polygon_to_bbox(value: object) -> tuple[float, float, float, float] | None:
    if not isinstance(value, (list, tuple)):
        return None
    if len(value) == 4 and all(isinstance(item, (int, float)) for item in value):
        left, top, right, bottom = (float(item) for item in value)
        return left, top, right, bottom
    points = [
        point
        for point in value
        if isinstance(point, (list, tuple))
        and len(point) >= 2
        and isinstance(point[0], (int, float))
        and isinstance(point[1], (int, float))
    ]
    if not points:
        return None
    xs = [float(point[0]) for point in points]
    ys = [float(point[1]) for point in points]
    return min(xs), min(ys), max(xs), max(ys)


def normalize_local_ocr_page(
    raw_result: dict[str, object],
    *,
    source_page: int,
    page_size: tuple[float, float],
) -> StructuredPage:
    """Normalize Paddle-style local OCR output without leaking SDK details."""
    texts = raw_result.get("rec_texts", [])
    scores = raw_result.get("rec_scores", [])
    polygons = raw_result.get("rec_polys") or raw_result.get("rec_boxes") or []
    if not isinstance(texts, list):
        texts = []
    if not isinstance(scores, list):
        scores = []
    if not isinstance(polygons, list):
        polygons = []
    blocks: list[OcrBlock] = []
    for index, value in enumerate(texts):
        text = str(value).strip()
        if not text:
            continue
        confidence = scores[index] if index < len(scores) else None
        blocks.append(
            OcrBlock(
                text,
                _polygon_to_bbox(polygons[index]) if index < len(polygons) else None,
                float(confidence) if isinstance(confidence, (int, float)) else None,
            )
        )
    return StructuredPage(
        source_page=source_page,
        width=float(page_size[0]),
        height=float(page_size[1]),
        extraction_mode="local_ocr",
        blocks=tuple(blocks),
    )


def normalize_native_pdf_page(
    raw_blocks: list[object],
    *,
    source_page: int,
    page_size: tuple[float, float],
) -> StructuredPage:
    """Normalize PyMuPDF text blocks into the structured page contract."""
    blocks: list[OcrBlock] = []
    for raw_block in raw_blocks:
        if not isinstance(raw_block, (tuple, list)) or len(raw_block) < 5:
            continue
        left, top, right, bottom, value = raw_block[:5]
        if not all(isinstance(item, (int, float)) for item in (left, top, right, bottom)):
            continue
        text = str(value).strip()
        if text:
            blocks.append(
                OcrBlock(
                    text,
                    (float(left), float(top), float(right), float(bottom)),
                    None,
                )
            )
    return StructuredPage(
        source_page=source_page,
        width=float(page_size[0]),
        height=float(page_size[1]),
        extraction_mode="native_text",
        blocks=tuple(blocks),
    )


def parse_reading_source(path: str | Path, source_format: str) -> ParsedDocument:
    source = Path(path)
    parsers = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "txt": _parse_text,
        "markdown": _parse_markdown,
        "epub": _parse_epub,
    }
    parser = parsers.get(source_format)
    if parser is None:
        raise ValueError("unsupported reading format")
    parsed = parser(source)
    chapters = [chapter for chapter in parsed.chapters if chapter.plain_text.strip()]
    keys = {chapter.chapter_key for chapter in chapters}
    outline = [item for item in parsed.outline if item.chapter_key in keys]
    return ParsedDocument(chapters, outline, parsed.assets)


def _safe_html(value: str) -> str:
    return nh3.clean(
        value,
        tags={"p", "h1", "h2", "h3", "h4", "blockquote", "strong", "em", "ul", "ol", "li", "code", "pre", "br", "hr", "img", "span"},
        attributes={
            "img": {"data-reading-asset", "alt", "title", "width", "height", "loading"},
            "span": {"class"},
        },
    )


def _plain_text_from_html(value: str, *, heading: str = "") -> str:
    soup = BeautifulSoup(value, "html.parser")
    if heading:
        first_heading = soup.find(["h1", "h2", "h3", "h4"])
        if first_heading and first_heading.get_text("", strip=True) == heading.strip():
            first_heading.decompose()
    for line_break in soup.find_all("br"):
        line_break.replace_with("\n")
    for block in soup.find_all(
        ["h1", "h2", "h3", "h4", "p", "li", "blockquote", "pre"]
    ):
        block.append("\n")
    lines = [
        re.sub(r"[ \t\r\f\v]+", " ", line).strip()
        for line in soup.get_text("", strip=False).splitlines()
    ]
    return "\n".join(line for line in lines if line)


def _paragraph_html(text: str) -> str:
    parts = [line.strip() for line in text.splitlines() if line.strip()]
    return "".join(f"<p>{html.escape(line)}</p>" for line in parts)


_PDF_LIST_ITEM = re.compile(r"^\s*(?:(\d+)[.)、]|([-•●▪]))\s*(.+?)\s*$")
_PDF_PARAGRAPH_END = re.compile(r"[。！？!?；;：:]\s*$")
_PDF_HEADING_END = re.compile(r"[，,、。！？!?；;：:]\s*$")
_PDF_ISOLATED_MARKER = re.compile(r"^(?:\d{1,3}|[-•●▪·])$")
_SCANNED_PDF_PERIOD = re.compile(r"^第[一二三四五六七八九十]+次国内革命战争时期$")
_SCANNED_PDF_DATE = re.compile(r"[（(]一九[〇零一二三四五六七八九十0-9]")
_SCANNED_PDF_LEADER = re.compile(r"[.…·]+\s*\d*(?:\s*[-—–]\s*\d+)?\s*$")


def _join_pdf_fragments(parts: list[str]) -> str:
    value = ""
    for part in parts:
        fragment = part.strip()
        if not fragment:
            continue
        needs_space = bool(
            value
            and re.search(r"[A-Za-z0-9)]$", value)
            and re.match(r"[A-Za-z0-9(]", fragment)
        )
        value += (" " if needs_space else "") + fragment
    return value


def pdf_text_to_html(
    text: str,
    *,
    promote_leading_headings: bool = False,
) -> str:
    """Rebuild semantic paragraphs from visual lines extracted from a PDF."""
    blocks: list[tuple[str, str | list[str]]] = []
    lines = text.splitlines()
    if promote_leading_headings:
        cursor = 0
        for kind, max_length in (("h1", 40), ("h2", 24)):
            while cursor < len(lines) and not lines[cursor].strip():
                cursor += 1
            if cursor >= len(lines):
                break
            candidate = lines[cursor].strip()
            if (
                len(candidate) > max_length
                or _PDF_HEADING_END.search(candidate)
                or _PDF_ISOLATED_MARKER.fullmatch(candidate)
                or _PDF_LIST_ITEM.match(candidate)
            ):
                break
            blocks.append((kind, candidate))
            cursor += 1
        lines = lines[cursor:]

    paragraph: list[str] = []
    list_kind: str | None = None
    list_items: list[str] = []
    current_item: list[str] = []

    def flush_paragraph() -> None:
        if paragraph:
            blocks.append(("p", _join_pdf_fragments(paragraph)))
            paragraph.clear()

    def flush_item() -> None:
        if current_item:
            list_items.append(_join_pdf_fragments(current_item))
            current_item.clear()

    def flush_list() -> None:
        nonlocal list_kind
        flush_item()
        if list_kind and list_items:
            blocks.append((list_kind, list_items.copy()))
        list_items.clear()
        list_kind = None

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            flush_list()
            continue
        if _PDF_ISOLATED_MARKER.fullmatch(line):
            continue

        item_match = _PDF_LIST_ITEM.match(line)
        if item_match:
            kind = "ol" if item_match.group(1) else "ul"
            flush_paragraph()
            if list_kind and list_kind != kind:
                flush_list()
            if list_kind is None:
                list_kind = kind
            flush_item()
            current_item.append(item_match.group(3))
            continue

        if list_kind:
            if current_item and _PDF_PARAGRAPH_END.search(current_item[-1]):
                flush_list()
            else:
                current_item.append(line)
                continue

        paragraph.append(line)
        if _PDF_PARAGRAPH_END.search(line):
            flush_paragraph()

    flush_paragraph()
    flush_list()

    rendered: list[str] = []
    for kind, content in blocks:
        if kind in {"p", "h1", "h2"}:
            rendered.append(f"<{kind}>{html.escape(str(content))}</{kind}>")
            continue
        items = "".join(f"<li>{html.escape(item)}</li>" for item in content)
        rendered.append(f"<{kind}>{items}</{kind}>")
    return "".join(rendered)


def _with_outline(
    chapters: list[ParsedChapter], assets: list[ParsedAsset] | None = None
) -> ParsedDocument:
    normalized: list[ParsedChapter] = []
    outline: list[ParsedOutlineItem] = []
    stack: list[tuple[int, str]] = []
    for position, chapter in enumerate(chapters):
        key = chapter.chapter_key or f"chapter-{position}"
        level = max(1, min(6, int(chapter.level or 1)))
        while stack and stack[-1][0] >= level:
            stack.pop()
        parent_key = chapter.parent_key or (stack[-1][1] if stack else None)
        value = ParsedChapter(
            chapter.title,
            chapter.plain_text,
            chapter.sanitized_html,
            chapter.source_page,
            key,
            level,
            parent_key,
            chapter.epub_cfi,
        )
        normalized.append(value)
        outline.append(
            ParsedOutlineItem(
                key,
                value.title,
                level,
                parent_key,
                value.source_page,
                value.epub_cfi,
            )
        )
        stack.append((level, key))
    unique_assets = {item.safe_name: item for item in (assets or [])}
    return ParsedDocument(normalized, outline, list(unique_assets.values()))


def _normalize_scanned_pdf_heading(value: str) -> str:
    return re.sub(r"[^\w\u4e00-\u9fff]", "", value)


def _clean_scanned_pdf_catalog_title(value: str) -> str:
    opening_positions = [
        position for mark in ("（", "(") if (position := value.find(mark)) >= 0
    ]
    if opening_positions and not re.search(r"[）)]", value[min(opening_positions) :]):
        return value[: min(opening_positions)].strip()
    return value


def _find_scanned_pdf_heading_target(
    heading: str, chapters: list[ParsedChapter]
) -> ParsedChapter | None:
    leading = re.split(r"[（(]", heading, maxsplit=1)[0]
    needle = _normalize_scanned_pdf_heading(leading)
    if len(needle) < 3:
        return None
    for chapter in chapters:
        heading_lines = [
            _normalize_scanned_pdf_heading(line)
            for line in chapter.plain_text.splitlines()[:5]
        ]
        if needle in heading_lines:
            return chapter
    return None


def build_scanned_pdf_outline(
    chapters: list[ParsedChapter],
) -> list[ParsedOutlineItem]:
    """Build navigation from verified catalog entries, not book-specific patterns."""
    ordered = sorted(
        (chapter for chapter in chapters if chapter.source_page is not None),
        key=lambda chapter: chapter.source_page or 0,
    )
    start_index = next(
        (
            index
            for index, chapter in enumerate(ordered)
            if re.search(r"目\s*录", chapter.plain_text)
        ),
        None,
    )
    if start_index is None:
        return []

    def has_catalog_entry(chapter: ParsedChapter) -> bool:
        return any(
            re.search(r"[.…·]{2,}\s*(?:[0-9IVXLCDM]+(?:\s*[-—–]\s*[0-9IVXLCDM]+)?)?\s*$", line)
            for line in chapter.plain_text.splitlines()
        )

    catalog_pages: list[ParsedChapter] = []
    for index, chapter in enumerate(ordered[start_index:], start=start_index):
        if index == start_index or has_catalog_entry(chapter):
            catalog_pages.append(chapter)
            continue
        break
    catalog_page_numbers = {
        chapter.source_page for chapter in catalog_pages if chapter.source_page is not None
    }
    content_pages = [
        chapter
        for chapter in ordered
        if chapter.source_page not in catalog_page_numbers
    ]
    outline: list[ParsedOutlineItem] = []
    seen_titles: set[str] = set()
    current_parent_key: str | None = None
    for chapter in catalog_pages:
        previous_title_fragment: str | None = None
        lines = [line.strip() for line in chapter.plain_text.splitlines() if line.strip()]
        for index, raw_line in enumerate(lines):
            leader_match = re.search(r"[.…·]{2,}\s*(?:[0-9IVXLCDM]+(?:\s*[-—–]\s*[0-9IVXLCDM]+)?)?\s*$", raw_line)
            raw_title = raw_line[: leader_match.start()].strip() if leader_match else raw_line
            article_lead = re.split(r"[（(]", raw_title, maxsplit=1)[0].strip()
            if leader_match and previous_title_fragment and article_lead.startswith(("而", "和", "的")):
                raw_title = previous_title_fragment + raw_title
                previous_title_fragment = None
            title = _clean_scanned_pdf_catalog_title(raw_title)
            if (
                not title
                or title in {"目", "录"}
                or _PDF_ISOLATED_MARKER.fullmatch(title)
                or title.startswith("毛泽东选集")
            ):
                continue
            next_has_leader = (
                index + 1 < len(lines)
                and bool(re.search(r"[.…·]{2,}", lines[index + 1]))
            )
            next_title = (
                re.split(r"[.…·]{2,}", lines[index + 1], maxsplit=1)[0].strip()
                if next_has_leader
                else ""
            )
            if next_has_leader and next_title.startswith(("而", "和", "的")):
                previous_title_fragment = raw_title
                continue
            is_group = not leader_match and next_has_leader
            is_truncated_entry = not leader_match and bool(re.search(r"[（(]", raw_title))
            if not leader_match and not is_group and not is_truncated_entry:
                previous_title_fragment = raw_title
                continue
            previous_title_fragment = None
            if title in seen_titles:
                continue
            target = _find_scanned_pdf_heading_target(title, content_pages)
            if target is None:
                continue
            seen_titles.add(title)
            if is_group:
                current_parent_key = target.chapter_key
                outline.append(
                    ParsedOutlineItem(
                        target.chapter_key,
                        title,
                        1,
                        source_page=target.source_page,
                    )
                )
                continue
            outline.append(
                ParsedOutlineItem(
                    target.chapter_key,
                    title,
                    2 if current_parent_key else 1,
                    current_parent_key,
                    target.source_page,
                )
            )
    return outline


def _parse_text(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8-sig").strip()
    chapter_pattern = re.compile(r"^\s*(第[一二三四五六七八九十零百千万0-9]+[章回节卷部篇]|Chapter\s+\d+).*$", re.MULTILINE)
    matches = list(chapter_pattern.finditer(text))
    
    if not matches:
        return _with_outline([ParsedChapter("全文", text, _paragraph_html(text))])
        
    chapters: list[ParsedChapter] = []
    if text[: matches[0].start()].strip():
        intro = text[: matches[0].start()].strip()
        chapters.append(ParsedChapter("前言", intro, _paragraph_html(intro)))
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        section = text[match.start():end].strip()
        title = match.group(0).strip()[:200]
        chapters.append(ParsedChapter(title, section, _paragraph_html(section)))
    return _with_outline(chapters)


_MARKDOWN_IMAGE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_DATA_IMAGE = re.compile(
    r"^data:(image/(?:png|jpeg|gif|webp|avif));base64,([A-Za-z0-9+/=\s]+)$",
    re.IGNORECASE,
)


def _prepare_markdown_images(
    text: str, *, local_root: Path | None = None
) -> tuple[str, list[ParsedAsset]]:
    assets: list[ParsedAsset] = []

    def replace(match: re.Match[str]) -> str:
        alt = html.escape(match.group(1), quote=True)
        target = match.group(2).strip().strip("<>")
        data: bytes | None = None
        mime_type = ""
        data_match = _DATA_IMAGE.fullmatch(target)
        try:
            if data_match:
                mime_type = data_match.group(1).lower()
                data = base64.b64decode(data_match.group(2), validate=True)
            elif local_root is not None:
                candidate = Path(target)
                candidate = candidate if candidate.is_absolute() else local_root / candidate
                resolved = candidate.resolve(strict=True)
                resolved.relative_to(local_root.resolve())
                data = resolved.read_bytes()
                mime_type = mimetypes.guess_type(resolved.name)[0] or ""
            if data is not None:
                asset = ParsedAsset.from_bytes(data, mime_type, alt=match.group(1))
                assets.append(asset)
                return (
                    f'<img data-reading-asset="{asset.safe_name}" '
                    f'alt="{alt}" loading="lazy">'
                )
        except (OSError, ValueError, ReadingAssetError, binascii.Error):
            pass
        return (
            '<span class="reading-image-placeholder">'
            '图片资源未随文档导入</span>'
        )

    return _MARKDOWN_IMAGE.sub(replace, text), assets


def _markdown_heading_matches(text: str) -> list[re.Match[str]]:
    masked = re.sub(
        r"^```.*?^```\s*$|^~~~.*?^~~~\s*$",
        lambda match: re.sub(r"[^\n]", " ", match.group(0)),
        text,
        flags=re.MULTILINE | re.DOTALL,
    )
    heading = re.compile(r"^(#{1,3})\s+(.+?)\s*$", re.MULTILINE)
    return list(heading.finditer(masked))


def _parse_markdown_text(
    text: str, *, local_root: Path | None = None
) -> ParsedDocument:
    text, assets = _prepare_markdown_images(text, local_root=local_root)
    matches = _markdown_heading_matches(text)
    if not matches:
        rendered = markdown.markdown(text, extensions=["extra", "sane_lists"])
        sanitized = _safe_html(rendered)
        return _with_outline(
            [ParsedChapter("全文", _plain_text_from_html(sanitized), sanitized)], assets
        )

    chapters: list[ParsedChapter] = []
    if text[: matches[0].start()].strip():
        intro = text[: matches[0].start()].strip()
        chapters.append(ParsedChapter("前言", intro, _safe_html(markdown.markdown(intro, extensions=["extra"]))))
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        section = text[match.end():end].strip()
        title = match.group(2).strip()[:200]
        title = re.sub(r"\*\*|\*", "", title)
        body = section or title
        rendered = markdown.markdown(f"## {title}\n\n{section}", extensions=["extra", "sane_lists"])
        sanitized = _safe_html(rendered)
        chapters.append(
            ParsedChapter(
                title,
                _plain_text_from_html(sanitized, heading=title) or body,
                sanitized,
                chapter_key=f"heading-{index}",
                level=len(match.group(1)),
            )
        )
    return _with_outline(chapters, assets)


def _parse_markdown(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8-sig").strip()
    return _parse_markdown_text(text, local_root=path.parent)


def _parse_pdf(path: Path) -> ParsedDocument:
    import pymupdf4llm
    import pymupdf
    
    with tempfile.TemporaryDirectory() as temp_dir:
        md_text = pymupdf4llm.to_markdown(
            str(path), write_images=True, image_path=temp_dir
        )
        parsed = _parse_markdown_text(
            md_text.strip(), local_root=Path(temp_dir)
        )

    with pymupdf.open(str(path)) as document:
        native_toc = document.get_toc()
    if native_toc:
        chapter_by_title = {
            re.sub(r"\s+", "", chapter.title): chapter
            for chapter in parsed.chapters
        }
        outline: list[ParsedOutlineItem] = []
        stack: list[tuple[int, str]] = []
        for level, title, page, *_ in native_toc:
            chapter = chapter_by_title.get(re.sub(r"\s+", "", str(title)))
            if chapter is None:
                continue
            while stack and stack[-1][0] >= int(level):
                stack.pop()
            parent_key = stack[-1][1] if stack else None
            outline.append(
                ParsedOutlineItem(
                    chapter.chapter_key,
                    str(title)[:200],
                    int(level),
                    parent_key,
                    max(0, int(page) - 1),
                )
            )
            stack.append((int(level), chapter.chapter_key))
        if outline:
            parsed.outline = outline
    return parsed


def _ensure_local_ocr_runner_available() -> None:
    """Fail clearly before scheduling pages when the local OCR package is absent."""
    if importlib.util.find_spec("paddleocr") is None:
        raise LocalOcrUnavailable("local_ocr_not_installed")


def run_local_ocr_runner(
    image_path: Path,
    result_path: Path,
    *,
    runner=None,
) -> dict[str, object]:
    """Run OCR for one image in a fresh Python process and read its JSON result."""
    execute = runner or subprocess.run
    command = [
        sys.executable,
        str(Path(__file__).with_name("ocr_runner.py")),
        "--input",
        str(image_path),
        "--output",
        str(result_path),
    ]
    try:
        completed = execute(
            command,
            check=False,
            capture_output=True,
            text=True,
            timeout=_LOCAL_OCR_RUNNER_TIMEOUT_SECONDS,
        )
    except (OSError, subprocess.TimeoutExpired):
        return {}
    if completed.returncode != 0 or not result_path.is_file():
        return {}
    try:
        payload = json.loads(result_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}
    result = payload.get("res") if isinstance(payload, dict) else None
    return result if isinstance(result, dict) else {}


def parse_pdf_native_structure(path: str | Path) -> ParsedDocument:
    """Return structured pages only when every PDF page has usable native text."""
    import pymupdf

    source = Path(path)
    pages: list[StructuredPage] = []
    native_toc = []
    with pymupdf.open(str(source)) as document:
        native_toc = document.get_toc()
        for page_number, page in enumerate(document):
            structured = classify_page_blocks(
                normalize_native_pdf_page(
                    list(page.get_text("blocks")),
                    source_page=page_number,
                    page_size=(page.rect.width, page.rect.height),
                )
            )
            if sum(len(block.text) for block in structured.body_blocks) < 20:
                return ParsedDocument([], [], [])
            pages.append(structured)
    pages = classify_document_pages(pages)
    chapters = [
        ParsedChapter(
            title=f"第 {page.source_page + 1} 页",
            plain_text="\n".join(block.text for block in page.body_blocks),
            sanitized_html=pdf_text_to_html(
                "\n".join(block.text for block in page.body_blocks),
                promote_leading_headings=page.source_page == 0,
            ),
            source_page=page.source_page,
            chapter_key=f"page-{page.source_page}",
        )
        for page in pages
    ]
    parsed = _with_outline(chapters)
    parsed.pages = pages
    embedded_outline = build_native_pdf_outline(native_toc, page_count=len(pages))
    if embedded_outline:
        parsed.outline = embedded_outline
    else:
        catalog_outline = build_scanned_pdf_outline(parsed.chapters)
        if catalog_outline:
            parsed.outline = catalog_outline
    return parsed


def build_native_pdf_outline(
    toc: list[object], *, page_count: int
) -> list[ParsedOutlineItem]:
    """Convert a PDF's embedded TOC to page-backed reader navigation."""
    outline: list[ParsedOutlineItem] = []
    stack: list[tuple[int, str]] = []
    for entry in toc:
        if not isinstance(entry, (tuple, list)) or len(entry) < 3:
            continue
        level, title, page = entry[:3]
        if not isinstance(level, int) or not isinstance(page, int):
            continue
        source_page = page - 1
        cleaned_title = str(title).strip()[:200]
        if level < 1 or not cleaned_title or not 0 <= source_page < page_count:
            continue
        while stack and stack[-1][0] >= level:
            stack.pop()
        key = f"page-{source_page}"
        outline.append(
            ParsedOutlineItem(
                key,
                cleaned_title,
                level,
                stack[-1][1] if stack else None,
                source_page,
            )
        )
        stack.append((level, key))
    return outline


def parse_pdf_with_local_ocr(path: str | Path) -> ParsedDocument:
    """Extract a PDF locally, preferring positioned native text over OCR."""
    import pymupdf

    source = Path(path)
    chapters: list[ParsedChapter] = []
    pages: list[StructuredPage] = []
    with pymupdf.open(str(source)) as document:
        with tempfile.TemporaryDirectory(prefix="reading-ocr-") as temp_dir:
            image_dir = Path(temp_dir)
            structured_by_page: dict[int, StructuredPage] = {}
            pending: dict[
                int, tuple[Future[dict[str, object]], int, int]
            ] = {}
            workers = local_ocr_worker_count(len(document))
            ocr_checked = False

            def collect_page(page_number: int) -> None:
                future, width, height = pending.pop(page_number)
                try:
                    data = future.result()
                except Exception:
                    data = {}
                structured_by_page[page_number] = classify_page_blocks(
                    normalize_local_ocr_page(
                        data,
                        source_page=page_number,
                        page_size=(width, height),
                    )
                )

            with ThreadPoolExecutor(max_workers=workers) as executor:
                for page_number, page in enumerate(document):
                    native = normalize_native_pdf_page(
                        list(page.get_text("blocks")),
                        source_page=page_number,
                        page_size=(page.rect.width, page.rect.height),
                    )
                    native_text_length = sum(
                        len(block.text) for block in native.blocks if block.text.strip()
                    )
                    if native_text_length >= 20:
                        structured_by_page[page_number] = classify_page_blocks(native)
                        continue
                    if not ocr_checked:
                        _ensure_local_ocr_runner_available()
                        ocr_checked = True
                    image_path = image_dir / f"page-{page_number:04d}.png"
                    result_path = image_dir / f"page-{page_number:04d}.json"
                    render_scale = local_ocr_render_scale(
                        page.rect.width, page.rect.height
                    )
                    pixmap = page.get_pixmap(
                        matrix=pymupdf.Matrix(render_scale, render_scale),
                        alpha=False,
                    )
                    pixmap.save(image_path)
                    pending[page_number] = (
                        executor.submit(run_local_ocr_runner, image_path, result_path),
                        pixmap.width,
                        pixmap.height,
                    )
                    if len(pending) >= workers * 2:
                        collect_page(min(pending))
                for page_number in sorted(pending):
                    collect_page(page_number)
            pages = [structured_by_page[page_number] for page_number in sorted(structured_by_page)]
    pages = classify_document_pages(pages)
    chapters = []
    for page in pages:
        quality = score_page_quality(page)
        if quality.status == "attention":
            text = "本页本地 OCR 质量不足，请查看原版 PDF。"
            rendered = "<p>本页本地 OCR 质量不足，请查看原版 PDF。</p>"
        else:
            text = "\n".join(block.text for block in page.body_blocks)
            rendered = pdf_text_to_html(
                text, promote_leading_headings=page.source_page == 0
            )
        chapters.append(
            ParsedChapter(
                title=f"第 {page.source_page + 1} 页",
                plain_text=text,
                sanitized_html=rendered,
                source_page=page.source_page,
                chapter_key=f"page-{page.source_page}",
            )
        )
    parsed = _with_outline(chapters)
    parsed.pages = pages
    catalog_outline = build_scanned_pdf_outline(parsed.chapters)
    if catalog_outline:
        parsed.outline = catalog_outline
    return parsed


def _parse_docx(path: Path) -> ParsedDocument:
    import docx
    from docx.oxml.ns import qn

    document = docx.Document(str(path))
    sections: list[tuple[str, int, list[str]]] = []
    assets: list[ParsedAsset] = []
    current_title = "全文"
    current_level = 1
    current_lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        style_name = (paragraph.style.name or "").lower()
        style_id = (paragraph.style.style_id or "").lower()
        heading_match = re.search(r"heading\s*([1-3])", style_name + style_id)
        if heading_match and text:
            if current_lines:
                sections.append((current_title, current_level, current_lines))
            current_title, current_lines = text[:200], []
            current_level = int(heading_match.group(1))
            continue
        if text:
            current_lines.append(text)
        for blip in paragraph._p.xpath(".//a:blip"):
            rel_id = blip.get(qn("r:embed"))
            part = document.part.related_parts.get(rel_id)
            try:
                asset = ParsedAsset.from_bytes(
                    part.blob, part.content_type, alt="文档图片"
                )
            except (AttributeError, ReadingAssetError):
                current_lines.append("[图片暂不可用]")
            else:
                assets.append(asset)
                current_lines.append(
                    f'<img data-reading-asset="{asset.safe_name}" '
                    'alt="文档图片" loading="lazy">'
                )
    if current_lines:
        sections.append((current_title, current_level, current_lines))
    chapters = []
    for index, (title, level, lines) in enumerate(sections):
        rendered = "".join(
            line if line.startswith("<img ") else f"<p>{html.escape(line)}</p>"
            for line in lines
        )
        chapters.append(
            ParsedChapter(
                title,
                "\n".join(line for line in lines if not line.startswith("<img ")),
                _safe_html(rendered),
                chapter_key=f"docx-{index}",
                level=level,
            )
        )
    return _with_outline(chapters, assets)


def _parse_epub(path: Path) -> ParsedDocument:
    with zipfile.ZipFile(path) as archive:
        container = ElementTree.fromstring(archive.read("META-INF/container.xml"))
        rootfile = next(
            element.attrib["full-path"]
            for element in container.iter()
            if element.tag.endswith("rootfile")
        )
        package = ElementTree.fromstring(archive.read(rootfile))
        base = str(Path(rootfile).parent)
        if base == ".":
            base = ""
        manifest_items = {
            element.attrib.get("id", ""): dict(element.attrib)
            for element in package.iter()
            if element.tag.endswith("item")
        }
        manifest = {
            item_id: attributes.get("href", "")
            for item_id, attributes in manifest_items.items()
        }
        spine = [
            element.attrib.get("idref", "")
            for element in package.iter()
            if element.tag.endswith("itemref")
        ]
        chapters: list[ParsedChapter] = []
        chapter_by_member: dict[str, ParsedChapter] = {}
        assets: list[ParsedAsset] = []
        for position, item_id in enumerate(spine):
            href = manifest.get(item_id)
            if not href or "://" in href:
                continue
            member = str(Path(base) / href) if base else href
            soup = BeautifulSoup(archive.read(member), "html.parser")
            for tag in soup(["script", "style", "iframe", "object", "embed"]):
                tag.decompose()
            root = soup.body or soup
            member_dir = posixpath.dirname(member)
            for image in root.find_all("img"):
                src = str(image.get("src") or "")
                asset: ParsedAsset | None = None
                try:
                    image_member = posixpath.normpath(
                        posixpath.join(member_dir, src.split("#", 1)[0])
                    )
                    if image_member.startswith("../") or image_member.startswith("/"):
                        raise ValueError
                    data = archive.read(image_member)
                    mime = mimetypes.guess_type(image_member)[0] or ""
                    asset = ParsedAsset.from_bytes(
                        data, mime, alt=str(image.get("alt") or "")
                    )
                except (KeyError, OSError, ValueError, ReadingAssetError):
                    placeholder = soup.new_tag("span")
                    placeholder["class"] = "reading-image-placeholder"
                    placeholder.string = "图片暂不可用"
                    image.replace_with(placeholder)
                else:
                    assets.append(asset)
                    image.attrs = {
                        "data-reading-asset": asset.safe_name,
                        "alt": str(image.get("alt") or ""),
                        "loading": "lazy",
                    }
            title_tag = root.find(["h1", "h2", "h3"])
            title = title_tag.get_text(" ", strip=True) if title_tag else f"第 {position + 1} 章"
            text = root.get_text("\n", strip=True)
            chapter = ParsedChapter(
                    title[:200],
                    text,
                    _safe_html(str(root)),
                    chapter_key=f"epub-{position}",
                    epub_cfi=member,
            )
            chapters.append(chapter)
            chapter_by_member[posixpath.normpath(member)] = chapter

        parsed = _with_outline(chapters, assets)
        nav_item = next(
            (
                attributes
                for attributes in manifest_items.values()
                if "nav" in attributes.get("properties", "").split()
            ),
            None,
        )
        if not nav_item:
            return parsed
        nav_member = posixpath.normpath(
            posixpath.join(base, nav_item.get("href", ""))
        )
        try:
            nav_soup = BeautifulSoup(archive.read(nav_member), "html.parser")
        except KeyError:
            return parsed
        nav = nav_soup.find("nav", attrs={"epub:type": "toc"}) or nav_soup.find("nav")
        root_list = nav.find(["ol", "ul"], recursive=False) if nav else None
        if root_list is None:
            return parsed
        nav_dir = posixpath.dirname(nav_member)
        outline: list[ParsedOutlineItem] = []

        def add_nav_items(list_tag, level: int, parent_key: str | None = None) -> None:
            for item in list_tag.find_all("li", recursive=False):
                anchor = item.find("a", recursive=False) or item.find("a")
                current_parent = parent_key
                if anchor and anchor.get("href"):
                    target = posixpath.normpath(
                        posixpath.join(
                            nav_dir, str(anchor.get("href")).split("#", 1)[0]
                        )
                    )
                    chapter = chapter_by_member.get(target)
                    if chapter:
                        outline.append(
                            ParsedOutlineItem(
                                chapter.chapter_key,
                                anchor.get_text(" ", strip=True)[:200] or chapter.title,
                                level,
                                parent_key,
                                epub_cfi=str(anchor.get("href")),
                            )
                        )
                        current_parent = chapter.chapter_key
                nested = item.find(["ol", "ul"], recursive=False)
                if nested is not None:
                    add_nav_items(nested, level + 1, current_parent)

        add_nav_items(root_list, 1)
        if outline:
            parsed.outline = outline
        return parsed
